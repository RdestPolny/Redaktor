"""
Redaktor AI — Document Handler
Obsługa PDF, DOCX, DOC: ekstrakcja tekstu, obrazów, renderowanie stron.
"""

import io
import logging
from pathlib import Path
from typing import List, Dict, Optional
from dataclasses import dataclass, field

import fitz  # PyMuPDF

logger = logging.getLogger(__name__)

# Skala renderowania (2.0 = 144 DPI)
RENDER_SCALE = 2.0

# Opcjonalne importy
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import mammoth
    MAMMOTH_AVAILABLE = True
except ImportError:
    MAMMOTH_AVAILABLE = False


@dataclass
class PageContent:
    """Zawartość pojedynczej strony."""
    page_number: int       # 1-indexed
    text: str              # Wyekstrahowany tekst (zachowana kolejność)
    images: List[Dict] = field(default_factory=list)   # metadane obrazów
    has_images: bool = False
    estimated_columns: int = 1
    is_mostly_image: bool = False   # strona głównie graficzna


class DocumentHandler:
    """Obsługa różnych formatów dokumentów."""

    def __init__(self, file_source, filename: str):
        self.filename = filename
        self.file_type = self._detect_file_type(filename)
        self._document = None
        self._html_content = None

        self.is_path = isinstance(file_source, (str, Path))
        if self.is_path:
            self.file_path = str(file_source)
            self.file_bytes = None
        else:
            # BytesIO lub bytes
            if hasattr(file_source, 'read'):
                self.file_bytes = file_source.read()
            else:
                self.file_bytes = bytes(file_source)
            self.file_path = None

        self._load_document()

    def _detect_file_type(self, filename: str) -> str:
        ext = Path(filename).suffix.lower()
        if ext == '.pdf':
            return 'pdf'
        elif ext == '.docx':
            if not DOCX_AVAILABLE:
                raise ValueError("Format DOCX wymaga: pip install python-docx")
            return 'docx'
        elif ext == '.doc':
            if not MAMMOTH_AVAILABLE:
                raise ValueError("Format DOC wymaga: pip install mammoth")
            return 'doc'
        else:
            raise ValueError(f"Nieobsługiwany format: {ext}")

    def _load_document(self):
        if self.file_type == 'pdf':
            if self.is_path:
                self._document = fitz.open(self.file_path)
            else:
                self._document = fitz.open(stream=self.file_bytes, filetype="pdf")
        elif self.file_type == 'docx':
            if self.is_path:
                self._document = DocxDocument(self.file_path)
            else:
                self._document = DocxDocument(io.BytesIO(self.file_bytes))
        elif self.file_type == 'doc':
            if self.is_path:
                with open(self.file_path, 'rb') as f:
                    result = mammoth.convert_to_html(f)
            else:
                result = mammoth.convert_to_html(io.BytesIO(self.file_bytes))
            self._html_content = result.value

    def get_page_count(self) -> int:
        if self.file_type == 'pdf':
            return len(self._document)
        elif self.file_type in ('docx', 'doc'):
            # Szacuj liczbę "stron" na podstawie słów
            if self.file_type == 'docx':
                text = '\n'.join(p.text for p in self._document.paragraphs)
            else:
                text = self._html_content
            words = len(text.split())
            return max(1, words // 400)
        return 0

    # ===== RENDEROWANIE STRONY =====

    def render_page_as_image(self, page_index: int) -> Optional[bytes]:
        """Renderuje stronę PDF jako PNG (zwraca bytes)."""
        if self.file_type != 'pdf' or not self._document:
            return None
        if page_index < 0 or page_index >= len(self._document):
            return None

        page = self._document[page_index]
        mat = fitz.Matrix(RENDER_SCALE, RENDER_SCALE)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        return pix.tobytes("png")

    # ===== EKSTRAKCJA TEKSTU =====

    def extract_page_content(self, page_index: int) -> PageContent:
        """Wyciąga tekst ze strony z zachowaniem kolejności czytania."""
        if self.file_type == 'pdf':
            return self._extract_pdf_page(page_index)
        elif self.file_type in ('docx', 'doc'):
            return self._extract_doc_page(page_index)
        return PageContent(page_number=page_index + 1, text="")

    def _extract_pdf_page(self, page_index: int) -> PageContent:
        """Ekstrakcja tekstu z PDF z uwzględnieniem wielokolumnowego układu."""
        page = self._document[page_index]
        page_rect = page.rect
        page_width = page_rect.width

        # Pobierz wszystkie bloki tekstu
        blocks = page.get_text("blocks", sort=True)
        # blocks: (x0, y0, x1, y1, text, block_no, block_type)
        # block_type: 0=text, 1=image

        text_blocks = [b for b in blocks if b[6] == 0 and b[4].strip()]
        image_blocks = [b for b in blocks if b[6] == 1]

        # Sprawdź czy strona jest głównie graficzna
        page_area = page_width * page_rect.height
        image_area = sum((b[2]-b[0]) * (b[3]-b[1]) for b in image_blocks)
        image_coverage = (image_area / page_area * 100) if page_area > 0 else 0
        is_mostly_image = image_coverage > 60 or (len(text_blocks) < 3 and len(image_blocks) > 0)

        # Detekcja kolumn — sprawdź czy bloki są wyraźnie podzielone w poziomie
        estimated_columns = self._estimate_columns(text_blocks, page_width)

        if estimated_columns >= 2:
            text = self._extract_multicolumn(text_blocks, page_width, estimated_columns)
        else:
            # Jedna kolumna — złącz w kolejności Y
            text = "\n\n".join(b[4].strip() for b in text_blocks)

        # Wyczyść artefakty
        text = self._clean_text(text)

        # Zbierz metadane obrazów
        image_meta = []
        for img in page.get_images(full=True):
            xref = img[0]
            try:
                info = self._document.extract_image(xref)
                if info["width"] >= 50 and info["height"] >= 50:
                    image_meta.append({
                        "xref": xref,
                        "width": info["width"],
                        "height": info["height"],
                        "ext": info["ext"],
                    })
            except Exception:
                pass

        return PageContent(
            page_number=page_index + 1,
            text=text,
            images=image_meta,
            has_images=len(image_meta) > 0,
            estimated_columns=estimated_columns,
            is_mostly_image=is_mostly_image,
        )

    def _estimate_columns(self, text_blocks: list, page_width: float) -> int:
        """Szacuje liczbę kolumn na podstawie rozkładu X bloków tekstu."""
        if len(text_blocks) < 4:
            return 1

        # Zbierz środki X bloków
        centers = [(b[0] + b[2]) / 2 for b in text_blocks]
        mid = page_width / 2

        # Policz bloki po lewej i prawej stronie środka
        left = sum(1 for c in centers if c < mid * 0.9)
        right = sum(1 for c in centers if c > mid * 1.1)

        if left >= 2 and right >= 2:
            # Sprawdź czy jest 3. kolumna
            third = page_width * 0.66
            far_right = sum(1 for c in centers if c > third)
            if far_right >= 2 and left >= 2:
                return 3
            return 2
        return 1

    def _extract_multicolumn(self, text_blocks: list, page_width: float, n_cols: int) -> str:
        """Ekstrahuje tekst z układu wielokolumnowego — kolumna po kolumnie."""
        col_width = page_width / n_cols
        columns: Dict[int, list] = {i: [] for i in range(n_cols)}

        for block in text_blocks:
            x_center = (block[0] + block[2]) / 2
            col_idx = min(int(x_center / col_width), n_cols - 1)
            columns[col_idx].append(block)

        # Posortuj każdą kolumnę od góry do dołu
        result_parts = []
        for col_idx in range(n_cols):
            col_blocks = sorted(columns[col_idx], key=lambda b: b[1])
            col_text = "\n\n".join(b[4].strip() for b in col_blocks)
            if col_text.strip():
                result_parts.append(col_text)

        return "\n\n".join(result_parts)

    def _clean_text(self, text: str) -> str:
        """Usuwa typowe artefakty z ekstrakcji PDF."""
        import re
        lines = text.split('\n')
        clean = []
        for line in lines:
            stripped = line.strip()
            # Usuń izolowane numery stron (np. "42", "- 42 -")
            if re.match(r'^[-–\s]*\d{1,4}[-–\s]*$', stripped):
                continue
            # Usuń URL-e redakcyjne
            if re.match(r'^www\.[a-z]+\.[a-z]+', stripped, re.IGNORECASE):
                continue
            # Połącz łamanie wyrazów na końcu linii ("transforma-\ntorem" → "transformatorem")
            if clean and stripped and clean[-1].endswith('-'):
                clean[-1] = clean[-1][:-1] + stripped
                continue
            clean.append(line)
        return '\n'.join(clean)

    def _extract_doc_page(self, page_index: int) -> PageContent:
        """Ekstrakcja 'strony' z DOCX/DOC (wirtualne stronicowanie)."""
        if self.file_type == 'docx':
            all_paragraphs = [p.text for p in self._document.paragraphs if p.text.strip()]
            words_per_page = 400
            all_words = []
            for para in all_paragraphs:
                all_words.extend(para.split())
                all_words.append('\n')
            start = page_index * words_per_page
            end = start + words_per_page
            text = ' '.join(all_words[start:end])
        else:
            import re
            plain = re.sub(r'<[^>]+>', ' ', self._html_content)
            words = plain.split()
            start = page_index * 400
            text = ' '.join(words[start:start + 400])

        return PageContent(page_number=page_index + 1, text=text)

    # ===== EKSTRAKCJA OBRAZÓW =====

    def extract_page_images(self, page_index: int) -> List[Dict]:
        """Zwraca listę obrazów ze strony jako bytes (PNG).

        Każdy element: { 'bytes': bytes, 'width': int, 'height': int, 'ext': str }
        """
        if self.file_type != 'pdf' or not self._document:
            return []
        if page_index < 0 or page_index >= len(self._document):
            return []

        page = self._document[page_index]
        results = []

        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                base_img = self._document.extract_image(xref)
                width = base_img.get("width", 0)
                height = base_img.get("height", 0)

                # Pomiń miniaturki i ikonki
                if width < 80 or height < 80:
                    continue

                # Konwertuj do PNG jeśli potrzeba
                img_bytes = base_img["image"]
                ext = base_img.get("ext", "png").lower()

                if ext not in ("png", "jpeg", "jpg"):
                    # Użyj PyMuPDF do konwersji
                    pix = fitz.Pixmap(self._document, xref)
                    if pix.n > 4:  # CMYK
                        pix = fitz.Pixmap(fitz.csRGB, pix)
                    img_bytes = pix.tobytes("png")
                    ext = "png"

                results.append({
                    "bytes": img_bytes,
                    "width": width,
                    "height": height,
                    "ext": ext,
                })
            except Exception as e:
                logger.debug("Błąd ekstrakcji obrazu xref=%s: %s", xref, e)
                continue

        return results
